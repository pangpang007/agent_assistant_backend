"""
文本提取器：从不同类型的文件中提取纯文本。
每种文件类型对应一个提取方法。
"""

import csv
import io
from pathlib import Path
from typing import Optional

import structlog

logger = structlog.get_logger()

# 支持的文件类型及最大大小限制（bytes）
MAX_FILE_SIZES = {
    "pdf": 50 * 1024 * 1024,     # 50 MB
    "txt": 10 * 1024 * 1024,     # 10 MB
    "md": 10 * 1024 * 1024,      # 10 MB
    "csv": 20 * 1024 * 1024,     # 20 MB
    "docx": 30 * 1024 * 1024,    # 30 MB
}

SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".txt": "txt",
    ".md": "md",
    ".markdown": "md",
    ".csv": "csv",
    ".docx": "docx",
}


def detect_file_type(filename: str) -> Optional[str]:
    """
    根据文件扩展名检测文件类型。
    返回文件类型字符串或 None（不支持的类型）。
    """
    ext = Path(filename).suffix.lower()
    return SUPPORTED_EXTENSIONS.get(ext)


def validate_file_size(file_size: int, file_type: str) -> None:
    """
    校验文件大小是否在允许范围内。
    超出限制时抛出 ValueError。
    """
    max_size = MAX_FILE_SIZES.get(file_type, 10 * 1024 * 1024)
    if file_size > max_size:
        max_mb = max_size // (1024 * 1024)
        raise ValueError(f"文件大小超出限制，{file_type} 类型最大允许 {max_mb}MB")


def extract_text(file_path: str, file_type: str) -> str:
    """
    根据文件类型分发到对应的文本提取方法。

    Args:
        file_path: 文件的本地存储路径
        file_type: 文件类型 ("pdf", "txt", "md", "csv", "docx")

    Returns:
        提取出的纯文本字符串

    Raises:
        ValueError: 不支持的文件类型
        FileNotFoundError: 文件不存在
        Exception: 提取过程中的其他错误
    """
    extractors = {
        "pdf": _extract_pdf,
        "txt": _extract_txt,
        "md": _extract_txt,     # Markdown 按纯文本处理
        "csv": _extract_csv,
        "docx": _extract_docx,
    }

    extractor = extractors.get(file_type)
    if extractor is None:
        raise ValueError(f"不支持的文件类型: {file_type}")

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    logger.info("extracting_text", file_path=file_path, file_type=file_type)
    text = extractor(path)

    # 基础清洗
    text = _clean_text(text)

    if not text.strip():
        raise ValueError("文件内容为空或无法提取有效文本")

    logger.info("text_extracted", file_path=file_path, char_count=len(text))
    return text


def _extract_pdf(path: Path) -> str:
    """
    使用 PyMuPDF (fitz) 提取 PDF 文本。

    PyMuPDF 优势：
    - 速度快（C 语言底层）
    - 支持复杂 PDF 布局
    - 可处理多页文档
    """
    import fitz  # PyMuPDF

    text_parts = []
    doc = fitz.open(str(path))

    try:
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text("text")
            if page_text.strip():
                # 每页之间加分隔符
                text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
    finally:
        doc.close()

    return "\n\n".join(text_parts)


def _extract_txt(path: Path) -> str:
    """
    直接读取 TXT/Markdown 文件。
    尝试多种编码，优先 UTF-8。
    """
    encodings = ["utf-8", "gbk", "gb2312", "latin-1"]

    for encoding in encodings:
        try:
            return path.read_text(encoding=encoding)
        except (UnicodeDecodeError, UnicodeError):
            continue

    # 最后尝试二进制读取
    raise ValueError(f"无法解码文件，支持的编码: {encodings}")


def _extract_csv(path: Path) -> str:
    """
    CSV 文件逐行转文本。
    将每行数据转为 "列名: 值" 的格式，保留结构化信息。
    """
    text_parts = []
    encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
    content = None

    for encoding in encodings:
        try:
            content = path.read_text(encoding=encoding)
            break
        except (UnicodeDecodeError, UnicodeError):
            continue

    if content is None:
        raise ValueError("无法解码 CSV 文件")

    reader = csv.DictReader(io.StringIO(content))

    if not reader.fieldnames:
        raise ValueError("CSV 文件没有表头")

    for row_num, row in enumerate(reader, 1):
        row_text = f"Row {row_num}: " + ", ".join(
            f"{key}: {value}" for key, value in row.items() if value
        )
        text_parts.append(row_text)

    return "\n".join(text_parts)


def _extract_docx(path: Path) -> str:
    """
    使用 python-docx 提取 DOCX 文档文本。
    提取所有段落文本和表格文本。
    """
    from docx import Document

    doc = Document(str(path))
    text_parts = []

    # 提取段落
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # 提取表格
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                text_parts.append(row_text)

    return "\n\n".join(text_parts)


def _clean_text(text: str) -> str:
    """
    基础文本清洗：
    1. 去除连续多个空行（超过 2 个空行合并为 2 个）
    2. 去除行首行尾空白
    3. 替换特殊空白字符
    """
    import re

    # 替换特殊空白字符
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\t", " ")
    text = text.replace("\u00a0", " ")  # non-breaking space

    # 去除连续多行空白
    text = re.sub(r"\n{3,}", "\n\n", text)

    # 去除每行首尾空白
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)

    return text.strip()
